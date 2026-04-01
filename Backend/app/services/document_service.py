import io
import json
import os
import re
import uuid

import pdfplumber
from docx import Document as DocxDocument
from fastapi import UploadFile
from sqlalchemy.orm import Session

from ..database.models import Document
from ..services.parser_service import auto_process_document
from ..services.rag_service import store_chunks

UPLOAD_DIR = "uploads/documents"
os.makedirs(UPLOAD_DIR, exist_ok=True)


def save_file(file: UploadFile):
    ext = file.filename.split(".")[-1]
    unique_name = f"{uuid.uuid4()}.{ext}"
    file_path = os.path.join(UPLOAD_DIR, unique_name)

    with open(file_path, "wb") as f:
        f.write(file.file.read())

    return file_path


def extract_text_from_file(file: UploadFile):
    filename = file.filename.lower()

    file.file.seek(0)
    file_bytes = file.file.read()

    if filename.endswith(".txt"):
        text = file_bytes.decode("utf-8", errors="ignore")
    elif filename.endswith(".pdf"):
        text = ""

        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    elif filename.endswith(".docx"):
        doc = DocxDocument(io.BytesIO(file_bytes))
        text = "\n".join([para.text for para in doc.paragraphs])
    else:
        raise ValueError("Unsupported file type")

    text = re.sub(r"\n\s*\n", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)

    return text.strip()


def create_document(
    title: str,
    file: UploadFile,
    db: Session,
    user_id: int,
):
    file_path = save_file(file)
    extracted_text = extract_text_from_file(file)

    if not extracted_text.strip():
        raise ValueError("Could not extract text from file")

    new_doc = Document(
        user_id=user_id,
        title=title,
        content=extracted_text,
        file_path=file_path,
    )

    db.add(new_doc)
    db.commit()
    db.refresh(new_doc)

    store_chunks(extracted_text, new_doc.id, db)

    processed = auto_process_document(extracted_text)

    new_doc.doc_type = processed["doc_type"]
    new_doc.structured_data = json.dumps(processed["structured_data"])

    db.add(new_doc)
    db.commit()

    return new_doc